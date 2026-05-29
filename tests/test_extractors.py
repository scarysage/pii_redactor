"""
Round-trip tests for extractors.py.

We build fixtures programmatically (no real client data on disk) and feed
them through redact_file() into a tmp_path.

The .xlsx test covers the load-bearing behavior from CLAUDE.md:
  - leading zeros are preserved on the FLAGGED column header path (because
    we mask wholesale, but other text columns keep their stored string form)
  - free-text columns get per-cell PII detection.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from extractors import redact_file


# ---------------------------------------------------------------------------
# .txt
# ---------------------------------------------------------------------------

def test_redact_txt_roundtrip(tmp_path: Path):
    src = tmp_path / "in.txt"
    dst = tmp_path / "out.txt"
    src.write_text(
        "Jane Doe owes $200. SSN 456-78-9012. Email jane@example.com.",
        encoding="utf-8",
    )
    result = redact_file(src, dst)

    out = dst.read_text(encoding="utf-8")
    assert "456-78-9012" not in out
    assert "jane@example.com" not in out
    assert "<US_SSN>" in out
    assert "<EMAIL_ADDRESS>" in out
    assert any(f.entity_type == "US_SSN" for f in result.findings)


# ---------------------------------------------------------------------------
# .docx
# ---------------------------------------------------------------------------

def test_redact_docx_roundtrip(tmp_path: Path):
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Client: Jane Doe")
    doc.add_paragraph("SSN: 456-78-9012")
    doc.add_paragraph("Routing: 021000021, Account 1234567890")
    doc.add_paragraph("Normal sentence with no PII at all.")
    doc.save(str(src))

    result = redact_file(src, dst)

    out = Document(str(dst))
    full_text = "\n".join(p.text for p in out.paragraphs)
    assert "456-78-9012" not in full_text
    assert "Jane Doe" not in full_text  # PERSON should be redacted
    assert "<US_SSN>" in full_text
    assert "Normal sentence" in full_text  # safe lines untouched

    # At least one redacted run should be styled bold + red.
    bold_red_seen = False
    for p in out.paragraphs:
        for run in p.runs:
            if run.bold and run.font.color and run.font.color.rgb is not None:
                # python-docx returns RGBColor; we check the hex equals C00000.
                if str(run.font.color.rgb) == "C00000":
                    bold_red_seen = True
    assert bold_red_seen, "Expected at least one bold+red redacted run"


# ---------------------------------------------------------------------------
# .xlsx
# ---------------------------------------------------------------------------

def test_redact_xlsx_header_flagged_column_masked_wholesale(tmp_path: Path):
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Clients"
    ws.append(["Client Name", "SSN", "Notes"])
    ws.append(["Jane Doe", "456-78-9012", "Routine filing"])
    ws.append(["John Smith", "555-12-3456", "Owes 2023 1099 to vendor"])
    wb.save(str(src))

    result = redact_file(src, dst)

    out_wb = load_workbook(str(dst))
    out_ws = out_wb["Clients"]

    # SSN column (col 2) should be masked wholesale.
    assert out_ws.cell(row=2, column=2).value == "<US_SSN>"
    assert out_ws.cell(row=3, column=2).value == "<US_SSN>"
    # Client Name column (col 1) should also be flagged by header.
    assert "<" in str(out_ws.cell(row=2, column=1).value)
    # The flagged cell should be set to text format so the tag is not re-parsed.
    assert out_ws.cell(row=2, column=2).number_format == "@"


def test_redact_xlsx_freetext_column_cell_level(tmp_path: Path):
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Notes"
    ws.append(["Date", "Notes"])
    ws.append(["2024-01-01", "Spoke with Jane Doe re EIN 12-3456789."])
    ws.append(["2024-01-02", "No PII here. Just a normal note."])
    wb.save(str(src))

    redact_file(src, dst)

    out_ws = load_workbook(str(dst))["Notes"]
    cell_b2 = str(out_ws.cell(row=2, column=2).value)
    assert "12-3456789" not in cell_b2
    assert "<US_EIN>" in cell_b2
    # The harmless row stays put.
    assert "normal note" in str(out_ws.cell(row=3, column=2).value)


def test_redact_xlsx_preserves_leading_zeros_in_non_flagged_text_column(
    tmp_path: Path,
):
    """
    Per CLAUDE.md: "Preserve leading zeros / text-formatted numbers".

    A column that is NOT flagged sensitive and contains text-formatted values
    like account-style strings ("0123") must not have its leading zeros stripped
    by our output pass. We only round-trip read -> write here; the redactor
    should not even touch this column.
    """
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Codes"
    ws.append(["Date", "Code"])
    # Force the Code column to text format so "0123" stays as a string.
    ws.cell(row=1, column=2).number_format = "@"
    ws.append(["2024-01-01", "0123"])
    ws.append(["2024-01-02", "00042"])
    wb.save(str(src))

    redact_file(src, dst)

    out_ws = load_workbook(str(dst))["Codes"]
    # Both values: leading zeros intact (no int coercion, no truncation).
    # These specific values don't trip any recognizer at the cell level either,
    # so they should survive untouched end-to-end.
    assert str(out_ws.cell(row=2, column=2).value) == "0123"
    assert str(out_ws.cell(row=3, column=2).value) == "00042"


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

def test_unsupported_extension_raises(tmp_path: Path):
    src = tmp_path / "weird.zzz"
    src.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError):
        redact_file(src, tmp_path / "out.zzz")
