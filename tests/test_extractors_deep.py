# Copyright (c) 2026 Vincent Shahinllari. All rights reserved.
"""
Deeper coverage for extractors.py beyond the basic round-trips.

Focus areas:
  * DOCX tables (PII inside table cells)
  * XLSX multi-sheet, numeric cells, empty rows/cells, formulas
  * extract_text() helper for each format
  * Unsupported file types raise the right error
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from extractors import extract_text, redact_file


# ---------------------------------------------------------------------------
# DOCX: tables
# ---------------------------------------------------------------------------

def test_docx_table_cells_get_redacted(tmp_path: Path):
    """PII inside table cells must be caught, not just paragraph text."""
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Roster:")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "SSN"
    table.rows[1].cells[0].text = "Jane Doe"
    table.rows[1].cells[1].text = "456-78-9012"
    doc.save(str(src))

    redact_file(src, dst)

    out = Document(str(dst))
    # Pull all text from all cells.
    cell_text: list[str] = []
    for t in out.tables:
        for row in t.rows:
            for cell in row.cells:
                cell_text.append(cell.text)
    joined = "\n".join(cell_text)
    assert "456-78-9012" not in joined
    assert "<US_SSN>" in joined


def test_docx_paragraphs_with_no_pii_untouched(tmp_path: Path):
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Quarter results were positive.")
    doc.add_paragraph("")  # empty paragraph
    doc.add_paragraph("No PII in this sentence either.")
    doc.save(str(src))

    redact_file(src, dst)

    out_text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    assert "Quarter results were positive." in out_text
    assert "No PII in this sentence either." in out_text


def test_docx_table_column_masking_catches_bare_numbers(tmp_path: Path):
    """The real-world regression from the engagement-letter test: account
    and routing numbers sit in cells with no surrounding context words, so
    the per-paragraph scan misses them or mis-tags them. The header-keyword
    column pass MUST mask the whole column even when individual cells
    don't have the right context."""
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Accounts on File")
    table = doc.add_table(rows=3, cols=4)
    table.rows[0].cells[0].text = "Client"
    table.rows[0].cells[1].text = "Bank"
    table.rows[0].cells[2].text = "Account #"
    table.rows[0].cells[3].text = "Routing #"

    table.rows[1].cells[0].text = "Aaron Strassler"
    table.rows[1].cells[1].text = "First Meridian Bank"
    table.rows[1].cells[2].text = "0048291756"
    table.rows[1].cells[3].text = "021000089"

    table.rows[2].cells[0].text = "Miriam Herbstman"
    table.rows[2].cells[1].text = "Northgate Securities"
    table.rows[2].cells[2].text = "739104825"
    table.rows[2].cells[3].text = "121000358"

    doc.save(str(src))
    redact_file(src, dst)

    out = Document(str(dst))
    cells_by_col: list[list[str]] = [[], [], [], []]
    for t in out.tables:
        for row in t.rows:
            for col_idx, cell in enumerate(row.cells):
                cells_by_col[col_idx].append(cell.text)

    # Header row stays intact.
    assert cells_by_col[2][0] == "Account #"
    assert cells_by_col[3][0] == "Routing #"

    # Account # column: data cells must be replaced with the account tag.
    # Critically, the raw digits MUST be gone -- this is what was failing.
    for raw in ("0048291756", "739104825"):
        for col_cells in cells_by_col:
            for cell_text in col_cells[1:]:  # skip header
                assert raw not in cell_text, (
                    f"raw account # {raw!r} leaked into output"
                )
    # The data rows in columns 2 and 3 should be tagged.
    for row_idx in (1, 2):
        assert "<" in cells_by_col[2][row_idx], (
            f"Account col row {row_idx} not masked: {cells_by_col[2][row_idx]!r}"
        )
        assert "<" in cells_by_col[3][row_idx], (
            f"Routing col row {row_idx} not masked: {cells_by_col[3][row_idx]!r}"
        )
    # And the raw routing values are gone too.
    for raw in ("021000089", "121000358"):
        for col_cells in cells_by_col:
            for cell_text in col_cells[1:]:
                assert raw not in cell_text, (
                    f"raw routing # {raw!r} leaked into output"
                )


def test_docx_table_without_sensitive_headers_uses_cell_level(tmp_path: Path):
    """A table whose row-1 headers don't match sensitive keywords must NOT
    get the wholesale-column treatment. Per-paragraph scan still runs."""
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Description"
    table.rows[1].cells[0].text = "Pencils"
    table.rows[1].cells[1].text = "A normal description"
    doc.save(str(src))

    redact_file(src, dst)

    out = Document(str(dst))
    cells = []
    for t in out.tables:
        for row in t.rows:
            for cell in row.cells:
                cells.append(cell.text)
    # Nothing should have been masked wholesale -- no `<` in any cell.
    assert "Pencils" in cells
    assert "A normal description" in cells


def test_docx_table_freetext_column_is_exempt(tmp_path: Path):
    """A column named Notes/Comments/Description must NOT be masked wholesale;
    cell-level redaction applies instead."""
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Date"
    table.rows[0].cells[1].text = "Notes"
    table.rows[1].cells[0].text = "2024-01-01"
    table.rows[1].cells[1].text = "Routine call, EIN on file: 12-3456789"
    doc.save(str(src))

    redact_file(src, dst)

    out = Document(str(dst))
    notes_cell = out.tables[0].rows[1].cells[1].text
    # The whole cell should NOT be replaced by a single tag.
    assert "Routine call" in notes_cell
    # But the EIN inside it should be redacted by the per-paragraph pass.
    assert "12-3456789" not in notes_cell
    assert "<US_EIN>" in notes_cell


# ---------------------------------------------------------------------------
# XLSX: multi-sheet
# ---------------------------------------------------------------------------

def test_xlsx_multi_sheet_each_processed(tmp_path: Path):
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    # Sheet 1: name + SSN headers
    ws1 = wb.active
    ws1.title = "Clients"
    ws1.append(["Client Name", "SSN"])
    ws1.append(["Jane Doe", "456-78-9012"])

    # Sheet 2: notes column (free-text)
    ws2 = wb.create_sheet("Notes")
    ws2.append(["Date", "Notes"])
    ws2.append(["2024-01-01", "Strassler called about EIN 12-3456789."])

    # Sheet 3: nothing sensitive
    ws3 = wb.create_sheet("Misc")
    ws3.append(["Code", "Count"])
    ws3.append(["A", 100])
    ws3.append(["B", 200])

    wb.save(str(src))
    redact_file(src, dst)

    out_wb = load_workbook(str(dst))
    # Sheet 1: SSN column wholesale
    assert out_wb["Clients"].cell(row=2, column=2).value == "<US_SSN>"
    # Sheet 2: cell-level redaction inside notes
    notes_cell = str(out_wb["Notes"].cell(row=2, column=2).value)
    assert "12-3456789" not in notes_cell
    assert "<US_EIN>" in notes_cell
    assert "Strassler" not in notes_cell
    # Sheet 3: untouched
    assert out_wb["Misc"].cell(row=2, column=1).value == "A"
    assert out_wb["Misc"].cell(row=2, column=2).value == 100


def test_xlsx_numeric_cell_preserved_when_no_pii(tmp_path: Path):
    """A pure-number cell with no PII characteristics must round-trip as a number."""
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Amounts"
    ws.append(["Item", "Cost"])
    ws.append(["Pencils", 12])
    ws.append(["Paper", 99.95])
    wb.save(str(src))

    redact_file(src, dst)

    out = load_workbook(str(dst))["Amounts"]
    assert out.cell(row=2, column=2).value == 12
    assert out.cell(row=3, column=2).value == 99.95


def test_xlsx_single_row_sheet(tmp_path: Path):
    """A sheet with only a header row must not crash."""
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Empty"
    ws.append(["Name", "SSN"])
    wb.save(str(src))

    # Should not raise.
    redact_file(src, dst)


def test_xlsx_completely_empty_sheet(tmp_path: Path):
    """A workbook with an empty default sheet must not crash."""
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    # Default sheet has no rows.
    wb.save(str(src))

    redact_file(src, dst)  # smoke


def test_xlsx_blank_cells_between_data(tmp_path: Path):
    """Blank cells/rows interleaved with data should be tolerated."""
    src = tmp_path / "in.xlsx"
    dst = tmp_path / "out.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Mixed"
    ws.append(["Name", "Notes"])
    ws.append(["Jane Doe", "First entry"])
    ws.append([None, None])  # blank row
    ws.append(["John Smith", None])  # half blank
    wb.save(str(src))

    redact_file(src, dst)

    out = load_workbook(str(dst))["Mixed"]
    # Column A flagged by header (Name) -> wholesale
    assert "<" in str(out.cell(row=2, column=1).value)
    assert "<" in str(out.cell(row=4, column=1).value)
    # Blank row stays blank.
    assert out.cell(row=3, column=1).value is None


# ---------------------------------------------------------------------------
# extract_text helper
# ---------------------------------------------------------------------------

def test_extract_text_txt(tmp_path: Path):
    p = tmp_path / "x.txt"
    p.write_text("hello world", encoding="utf-8")
    assert extract_text(p) == "hello world"


def test_extract_text_docx(tmp_path: Path):
    p = tmp_path / "x.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.save(str(p))
    text = extract_text(p)
    assert "alpha" in text
    assert "beta" in text


def test_extract_text_xlsx(tmp_path: Path):
    p = tmp_path / "x.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["a", "b"])
    ws.append(["c", "d"])
    wb.save(str(p))
    text = extract_text(p)
    assert "a" in text and "b" in text
    assert "c" in text and "d" in text


def test_extract_text_unsupported(tmp_path: Path):
    import pytest
    p = tmp_path / "weird.bin"
    p.write_bytes(b"binary")
    with pytest.raises(ValueError):
        extract_text(p)
