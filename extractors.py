# Copyright (c) 2026 Vincent Shahinllari. All rights reserved.
"""
File -> redacted file pipeline, per format.

The contract every extractor exposes:

    extract_text(src_path)               -> str
        For display / analysis when we just want the raw text.

    redact_file(src_path, dst_path)      -> ExtractionResult
        Build a redacted copy of the file at dst_path. Returns an
        ExtractionResult with all findings collected during the pass.

For .txt and .pdf we render output as a plain text file with `<TYPE>` tags
inline. For .docx and .xlsx we preserve the original document and overwrite
the offending runs/cells in place, so formatting survives.

Supported formats:
    .txt   -> plain text
    .pdf   -> text-only output (we do not rewrite PDFs in place; per CLAUDE.md,
              "keep PDF output simple")
    .docx  -> rewritten .docx with redacted runs styled bold + red
    .xlsx  -> rewritten .xlsx with hybrid column/cell redaction

Excel handling: see redact_xlsx() docstring -- this is the trickiest part.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import pdfplumber
from docx import Document
from docx.shared import RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill

from redactor import Finding, analyze, redact


# Style applied to redacted spans in DOCX/XLSX output. ONE consistent style
# everywhere (per CLAUDE.md): bold + red.
REDACT_COLOR_HEX = "C00000"  # deep red, readable on white
REDACT_RGB = RGBColor(0xC0, 0x00, 0x00)
XLSX_REDACT_FILL = PatternFill(
    start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"
)  # pale yellow background -- the bold+red text is the primary signal


# Headers that indicate a free-text column. These are explicitly NOT flagged for
# wholesale masking even if a sample-scan would otherwise trigger -- per CLAUDE.md
# "Run cell-level detection on free-text / notes columns." A note column may have
# incidental PII in some rows; we want to redact spans, not nuke the whole column.
FREETEXT_HEADER_KEYWORDS = [
    "note",
    "notes",
    "comment",
    "comments",
    "description",
    "memo",
    "remarks",
    "details",
]

# Headers that indicate a column should be redacted wholesale.
# Match is case-insensitive and substring-based (e.g. "Client SSN" matches "ssn").
SENSITIVE_HEADER_KEYWORDS = [
    "ssn",
    "social security",
    "ein",
    "tax id",
    "taxpayer id",
    "fein",
    "routing",
    "aba",
    "account number",
    "acct number",
    "acct #",
    "account #",
    "bank account",
    "client name",
    "customer name",
    "first name",
    "last name",
    "full name",
    "dob",
    "date of birth",
    "address",
    "email",
    "phone",
]


@dataclass
class ExtractionResult:
    """Returned by redact_file(). Carries findings for the review screen."""

    text: str  # the (extracted, post-redact) plain text view of the file
    findings: list[Finding] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)  # warnings / context for the UI


# ---------------------------------------------------------------------------
# Content-signature ("magic byte") check
# ---------------------------------------------------------------------------
# Defense in depth: the UI only accepts .txt/.pdf/.docx/.xlsx by extension,
# but an extension is trivial to fake. Before we hand a file to a parser
# (pdfminer / python-docx / openpyxl), confirm its first bytes actually match
# the claimed type. This narrows the chance that a file renamed to ".pdf" is
# fed to the PDF parser as a parser-exploit payload.
#
# Deliberately lenient so we never reject a genuine document:
#   * .pdf   -> "%PDF" must appear within the first 1024 bytes (the PDF spec
#               allows leading bytes before the header; some real-world PDFs
#               have a BOM or stray whitespace first).
#   * .docx  -> ZIP container, must start with the local-file-header magic
#     .xlsx    "PK\x03\x04" (Office Open XML files are ZIP archives).
#   * .txt   -> no check; any byte content is legitimate plain text.

def _verify_file_signature(src: Path) -> None:
    """Raise ValueError if `src`'s contents don't match its extension.

    Only checks formats with a well-defined signature. Unknown extensions are
    left to the dispatch in redact_file() to reject.
    """
    suffix = src.suffix.lower()
    try:
        head = src.read_bytes()[:1024]
    except OSError:
        # Let the actual parser surface the read error with its own message.
        return

    if suffix == ".pdf":
        if b"%PDF" not in head:
            raise ValueError("File is named .pdf but is not a PDF document.")
    elif suffix in (".docx", ".xlsx"):
        # Office Open XML is a ZIP archive. Empty/new archives use other PK
        # signatures, but a real document always starts with a local file
        # header, "PK\x03\x04".
        if not head.startswith(b"PK\x03\x04"):
            raise ValueError(
                f"File is named {suffix} but is not a valid Office document."
            )
    # .txt and anything else: no signature to check.


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

def redact_file(src: str | Path, dst: str | Path) -> ExtractionResult:
    """
    Dispatch on file extension. dst is the path to write the redacted output.
    """
    src = Path(src)
    dst = Path(dst)
    suffix = src.suffix.lower()
    # Confirm the bytes match the extension before any parser touches them.
    _verify_file_signature(src)
    if suffix == ".txt":
        return redact_txt(src, dst)
    if suffix == ".pdf":
        return redact_pdf(src, dst)
    if suffix == ".docx":
        return redact_docx(src, dst)
    if suffix == ".xlsx":
        return redact_xlsx(src, dst)
    raise ValueError(
        f"Unsupported file type: {suffix}. Supported: .txt .pdf .docx .xlsx"
    )


def extract_text(src: str | Path) -> str:
    """Best-effort text view of a file (no redaction). Used for the review UI."""
    src = Path(src)
    suffix = src.suffix.lower()
    if suffix == ".txt":
        return src.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return _pdf_to_text(src)
    if suffix == ".docx":
        return "\n".join(p.text for p in Document(str(src)).paragraphs)
    if suffix == ".xlsx":
        # Cheap view: tab-separated per row, blank line between sheets.
        wb = load_workbook(str(src), data_only=False)
        chunks: list[str] = []
        for ws in wb.worksheets:
            chunks.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                chunks.append("\t".join("" if c is None else str(c) for c in row))
            chunks.append("")
        return "\n".join(chunks)
    raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# .txt
# ---------------------------------------------------------------------------

def redact_txt(src: Path, dst: Path) -> ExtractionResult:
    text = src.read_text(encoding="utf-8", errors="replace")
    redacted, findings = redact(text)
    dst.write_text(redacted, encoding="utf-8")
    return ExtractionResult(text=redacted, findings=findings)


# ---------------------------------------------------------------------------
# .pdf
# ---------------------------------------------------------------------------

def _pdf_to_text(src: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(str(src)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n\n".join(parts)


def redact_pdf(src: Path, dst: Path) -> ExtractionResult:
    """
    PDFs are output as plain text with <TYPE> tags. We do NOT try to rewrite
    the PDF in place -- per CLAUDE.md, "keep PDF output simple". A real PDF
    overlay would need to handle layout, fonts, and redaction-of-redactions
    (visible black-bar style), which is out of scope for v1.

    dst should be a .txt path; if it is .pdf, we rewrite it to .txt and add a
    note so the UI can tell the user.
    """
    notes: list[str] = []
    text = _pdf_to_text(src)
    redacted, findings = redact(text)

    if dst.suffix.lower() == ".pdf":
        dst = dst.with_suffix(".txt")
        notes.append(
            "PDF input was redacted as a .txt file -- v1 does not rewrite PDFs in place."
        )

    dst.write_text(redacted, encoding="utf-8")
    return ExtractionResult(text=redacted, findings=findings, notes=notes)


# ---------------------------------------------------------------------------
# .docx
# ---------------------------------------------------------------------------

def redact_docx(src: Path, dst: Path) -> ExtractionResult:
    """
    Open the docx, walk every paragraph, replace PII spans in-place at the run
    level, and style the replacements bold + red. Tables get the same treatment.

    Why per-paragraph (not whole-document text):
        python-docx exposes text as runs inside paragraphs. We need to keep
        formatting. The simplest approach that preserves formatting is to
        analyze paragraph by paragraph (or cell by cell for tables), and
        rewrite the runs of that paragraph when there is a finding.
    """
    doc = Document(str(src))
    all_findings: list[Finding] = []

    def process(paragraph) -> None:
        text = paragraph.text
        if not text.strip():
            return
        findings = analyze(text)
        if not findings:
            return
        all_findings.extend(findings)
        _rewrite_paragraph(paragraph, text, findings)

    for p in doc.paragraphs:
        process(p)

    # Tables get the per-paragraph scan AND a column-header pass. The latter
    # mirrors the XLSX hybrid approach: in a real engagement letter or 1099
    # roster, sensitive values (SSN, Account #, Routing #) live in cells whose
    # only "context" is the header in row 1. Per-paragraph analysis can't see
    # that context, so a bare 9-digit cell gets mis-tagged as PHONE_NUMBER or
    # DATE_TIME -- or missed entirely. The header pass catches these.
    for table in doc.tables:
        _redact_docx_table(table, process, all_findings)

    doc.save(str(dst))

    # Build the plain-text view for the review screen by re-extracting after save.
    review_text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    return ExtractionResult(text=review_text, findings=all_findings)


def _redact_docx_table(table, process_paragraph, all_findings: list[Finding]) -> None:
    """Per-paragraph scan plus header-keyword column masking for a DOCX table.

    Sequence:
      1. Run the per-paragraph analyzer on every cell. Catches inline mentions
         (e.g. an email or PERSON name written inside a cell with surrounding
         context words).
      2. Treat row 1 as headers. If any header matches a SENSITIVE_HEADER
         keyword, mask every data cell in that column wholesale -- replacing
         the cell with the column's tag (e.g. `<US_BANK_ACCOUNT>`), styled
         bold + red.

    The masking pass overwrites whatever the per-paragraph scan produced in
    those cells. That's intentional: a header-flagged column gets the correct
    tag (US_BANK_ACCOUNT) even if per-paragraph analysis mis-tagged the value
    as PHONE_NUMBER for lack of context.

    Free-text columns (Notes / Comments / Memo / ...) are exempt from
    wholesale masking, exactly like in XLSX -- they keep the per-paragraph
    pass instead.
    """
    # Step 1: per-paragraph scan everywhere.
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)

    # Step 2: header-based column masking. Needs at least header + 1 data row.
    if len(table.rows) < 2:
        return

    header_cells = table.rows[0].cells
    headers_lower = [c.text.strip().lower() for c in header_cells]

    flagged_cols: dict[int, str] = {}
    for col_idx, h in enumerate(headers_lower):
        if not h:
            continue
        if any(kw in h for kw in FREETEXT_HEADER_KEYWORDS):
            continue
        for kw in SENSITIVE_HEADER_KEYWORDS:
            if kw in h:
                flagged_cols[col_idx] = _tag_for_header(h)
                break

    if not flagged_cols:
        return

    for row in table.rows[1:]:
        for col_idx, cell in enumerate(row.cells):
            if col_idx not in flagged_cols:
                continue
            if not cell.text.strip():
                continue
            tag = f"<{flagged_cols[col_idx]}>"
            _replace_cell_text(cell, tag)
            all_findings.append(Finding(
                start=0,
                end=len(tag),
                entity_type=flagged_cols[col_idx],
                score=1.0,
                text=cell.text,
            ))


def _replace_cell_text(cell, replacement: str) -> None:
    """Clear every paragraph and run in `cell`, leaving a single styled run.

    We can't just set cell.text = replacement because python-docx would
    insert it as a plain run with no styling. To get the bold + red style
    consistent with the rest of the output, we clear the existing structure
    and add one new run on a single paragraph.
    """
    # Drop every paragraph after the first; keep one to write into.
    paragraphs = cell.paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)

    p = cell.paragraphs[0]
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    run = p.add_run(replacement)
    run.bold = True
    run.font.color.rgb = REDACT_RGB


def _rewrite_paragraph(paragraph, text: str, findings: Iterable[Finding]) -> None:
    """
    Replace `paragraph` contents with the redacted version, styling each
    redacted span bold + red.

    We blow away the existing runs and recreate them. This loses inline
    formatting *inside* the paragraph (bold/italic that the user applied),
    but it's the only reliable way to inject styled redaction tags without a
    much more complex run-splitting routine. Acceptable trade-off for v1:
    redaction correctness > preserving inline formatting.
    """
    # Build (start, end, tag) segments in order.
    sorted_findings = sorted(findings, key=lambda f: (f.start, f.end))
    segments: list[tuple[str, bool]] = []  # (text, is_redacted)
    cursor = 0
    last_end = -1
    for f in sorted_findings:
        if f.start < last_end:
            # Overlapping: skip, the earlier (longer) match wins.
            continue
        if f.start > cursor:
            segments.append((text[cursor:f.start], False))
        segments.append((f"<{f.entity_type}>", True))
        cursor = f.end
        last_end = f.end
    if cursor < len(text):
        segments.append((text[cursor:], False))

    # Clear existing runs by removing the underlying XML children.
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    for seg_text, is_red in segments:
        run = paragraph.add_run(seg_text)
        if is_red:
            run.bold = True
            run.font.color.rgb = REDACT_RGB


# ---------------------------------------------------------------------------
# .xlsx
# ---------------------------------------------------------------------------

def redact_xlsx(src: Path, dst: Path) -> ExtractionResult:
    """
    Hybrid column + cell redaction, per CLAUDE.md "Excel handling".

    For each sheet, looping all sheets:

      1. Treat row 1 as the header row.
      2. Flag a column as "sensitive" if either:
           - its header matches one of SENSITIVE_HEADER_KEYWORDS, OR
           - a sample-scan of its non-header cells turns up PII findings
             in more than ~40% of non-empty cells.
         Flagged columns get masked wholesale: every non-empty data cell is
         replaced with a single `<COLUMN_REDACTED>` tag.
      3. For all other columns, run cell-level detection and replace any
         PII spans inline with `<TYPE>` tags.
      4. Style every modified cell bold + red on a pale-yellow fill so the
         review screen reads cleanly.

    Preserving leading zeros / text-formatted numbers:
        We open the workbook with data_only=False and read each cell's `.value`
        but stringify it ourselves -- we never re-cast to int. When we write
        back, we set number_format='@' (text) on any cell we redact so the
        replacement string is not re-interpreted as a number.
    """
    wb = load_workbook(str(src), data_only=False)
    all_findings: list[Finding] = []
    notes: list[str] = []
    review_chunks: list[str] = []

    for ws in wb.worksheets:
        review_chunks.append(f"# Sheet: {ws.title}")

        if ws.max_row < 1 or ws.max_column < 1:
            review_chunks.append("(empty sheet)")
            continue

        # Headers: row 1. We tolerate None / non-string headers.
        headers: list[str] = []
        for col_idx in range(1, ws.max_column + 1):
            h = ws.cell(row=1, column=col_idx).value
            headers.append("" if h is None else str(h))

        flagged_cols = _flag_sensitive_columns(ws, headers)

        # Walk data rows (skip header).
        for row_idx in range(2, ws.max_row + 1):
            row_chunks: list[str] = []
            for col_idx in range(1, ws.max_column + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                raw = cell.value
                if raw is None:
                    row_chunks.append("")
                    continue

                # Preserve the stored string form. For text cells this is the
                # value as the user typed it (leading zeros intact when the
                # cell's number_format is '@' / text).
                text = str(raw)

                if col_idx in flagged_cols and text.strip():
                    cell.value = f"<{flagged_cols[col_idx]}>"
                    cell.number_format = "@"
                    cell.font = Font(bold=True, color=REDACT_COLOR_HEX)
                    cell.fill = XLSX_REDACT_FILL
                    row_chunks.append(cell.value)
                    continue

                # Cell-level scan for non-flagged columns.
                findings = analyze(text)
                if not findings:
                    row_chunks.append(text)
                    continue

                all_findings.extend(findings)
                new_text = _replace_spans(text, findings)
                cell.value = new_text
                cell.number_format = "@"
                cell.font = Font(bold=True, color=REDACT_COLOR_HEX)
                cell.fill = XLSX_REDACT_FILL
                row_chunks.append(new_text)

            review_chunks.append("\t".join(row_chunks))

        if flagged_cols:
            pretty = ", ".join(
                f"{headers[c-1] or f'col{c}'} -> <{tag}>"
                for c, tag in flagged_cols.items()
            )
            notes.append(f"Sheet '{ws.title}': masked columns: {pretty}")

        review_chunks.append("")

    wb.save(str(dst))
    return ExtractionResult(
        text="\n".join(review_chunks), findings=all_findings, notes=notes
    )


def _flag_sensitive_columns(ws, headers: list[str]) -> dict[int, str]:
    """
    Return {1-based column index: redaction tag} for every column we should mask
    wholesale.

    Two ways a column gets flagged:
      (a) header keyword match -> tag is chosen from the keyword
      (b) sample scan of up to 20 non-empty cells finds PII in > 40% of them
          -> tag is "COLUMN_REDACTED" (we do not know the exact type)
    """
    flagged: dict[int, str] = {}

    for col_idx, header in enumerate(headers, start=1):
        h = header.lower()
        for kw in SENSITIVE_HEADER_KEYWORDS:
            if kw in h:
                flagged[col_idx] = _tag_for_header(h)
                break

    # Sample-scan: only for columns not already flagged and not explicitly
    # free-text. Free-text columns (Notes, Comments, ...) are handled at the
    # cell level so we preserve the surrounding sentence; see CLAUDE.md.
    #
    # Why MIN_HITS as well as HIT_RATIO: on a tiny sheet (say 2 rows), a single
    # false-positive (e.g. spaCy flagging "00099" as DATE_TIME) would clear a
    # 40% ratio and incorrectly mask the whole column. Requiring at least a
    # few hits avoids that on small inputs while still working on real-world
    # spreadsheets with dozens of rows.
    SAMPLE_LIMIT = 20
    HIT_RATIO = 0.4
    MIN_HITS = 3
    for col_idx, header in enumerate(headers, start=1):
        if col_idx in flagged:
            continue
        h = header.lower()
        if any(kw in h for kw in FREETEXT_HEADER_KEYWORDS):
            continue
        samples: list[str] = []
        for row_idx in range(2, ws.max_row + 1):
            if len(samples) >= SAMPLE_LIMIT:
                break
            v = ws.cell(row=row_idx, column=col_idx).value
            if v is None:
                continue
            s = str(v).strip()
            if s:
                samples.append(s)
        if not samples:
            continue
        hits = sum(1 for s in samples if analyze(s))
        if hits >= MIN_HITS and hits / len(samples) >= HIT_RATIO:
            flagged[col_idx] = "COLUMN_REDACTED"

    return flagged


# Sanity: every free-text keyword should be checked BEFORE the sample-scan loop,
# never as a "sensitive" header. Surface drift early if someone adds an overlap.
assert not (set(FREETEXT_HEADER_KEYWORDS) & set(SENSITIVE_HEADER_KEYWORDS)), (
    "FREETEXT and SENSITIVE header keyword lists overlap -- pick one per keyword."
)


# Mapping from header keyword -> entity tag used for column-wide masking.
# Keep this aligned with redactor.DEFAULT_ENTITIES.
_HEADER_TAGS = [
    ("ssn", "US_SSN"),
    ("social security", "US_SSN"),
    ("ein", "US_EIN"),
    ("tax id", "US_EIN"),
    ("taxpayer id", "US_EIN"),
    ("fein", "US_EIN"),
    ("routing", "US_BANK_ROUTING"),
    ("aba", "US_BANK_ROUTING"),
    ("account number", "US_BANK_ACCOUNT"),
    ("acct number", "US_BANK_ACCOUNT"),
    ("acct #", "US_BANK_ACCOUNT"),
    ("account #", "US_BANK_ACCOUNT"),
    ("bank account", "US_BANK_ACCOUNT"),
    ("client name", "PERSON"),
    ("customer name", "PERSON"),
    ("first name", "PERSON"),
    ("last name", "PERSON"),
    ("full name", "PERSON"),
    ("dob", "DATE_TIME"),
    ("date of birth", "DATE_TIME"),
    ("address", "LOCATION"),
    ("email", "EMAIL_ADDRESS"),
    ("phone", "PHONE_NUMBER"),
]


def _tag_for_header(header_lower: str) -> str:
    for kw, tag in _HEADER_TAGS:
        if kw in header_lower:
            return tag
    return "COLUMN_REDACTED"


def _replace_spans(text: str, findings: Iterable[Finding]) -> str:
    """Walk findings right-to-left, splice in `<TYPE>` tags. Skips overlaps."""
    out = text
    last_end = len(text) + 1
    for f in sorted(findings, key=lambda x: x.start, reverse=True):
        if f.end > last_end:
            continue
        out = out[: f.start] + f"<{f.entity_type}>" + out[f.end:]
        last_end = f.start
    return out
